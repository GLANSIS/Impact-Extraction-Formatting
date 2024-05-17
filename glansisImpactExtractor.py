#!/usr/bin/env python
# coding: utf-8

# In[ ]:

# Import libraries
import streamlit as st           # Create webpage
import pandas as pd              # Manage data tables
import requests                  # Pulls HTML code from webpage
from bs4 import BeautifulSoup    # HTML parsing
import re                        # Edit text strings
from docx import Document        # Create and edit Word Document
import base64                    # Encode data
from io import BytesIO           # Manage byte data
from selenium import webdriver                                   # Automate web browser interaction
from selenium.webdriver.chrome.options import Options            # Use'headless' browser options
from selenium.webdriver.common.by import By                      # Find HTML elements 
from selenium.webdriver.support.ui import Select                 # Use for dropdown selection
from selenium.webdriver.support.ui import WebDriverWait          # Command driver to wait 
from selenium.webdriver.support import expected_conditions as EC # Specify driver conditions for pause


# Function to scrape website
def scrape_impacts(num):
    
    # set URL
    url = 'https://nas.er.usgs.gov/queries/greatlakes/Impacts/ImpactsInfo.aspx?speciesID=' + num

    # Open web page using a headless Selenium webdriver
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    driver = webdriver.Chrome(options = chrome_options)
    driver.get(url)

    # Find dropdown elements on the web page and select 300 to view all impacts
    type_dropdown = Select(WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, 'body_ResultsPerPageDD'))))
    type_dropdown.select_by_visible_text(str(300))

    # Get HTML script and pull table from web page
    html = driver.page_source
    soup = BeautifulSoup(html,'html.parser')
    table = soup.find('table', {"id": "body_myGridView"})

    # Close the webdriver
    driver.quit()

    # Extract table information
    data = []
    for row in table.find_all('tr')[2:-2]:
        
        # Pull both headers and data cells
        cells = row.find_all(['th', 'td'])
        
        # Pull content from cells - include separator so italicize words don't lose space
        row_data = [cell.get_text(strip = True, separator = " ") for cell in cells]
        
        # Add to data list
        data.append(row_data)

    # Convert list to DataFrame
    impact_table = pd.DataFrame(data[1:], columns = data[0])

    # Filter rows based on user Impact ID input
    if first_impact_id != 'NA':
        selected_rows = impact_table[impact_table['Impact ID'] >= first_impact_id]
    else:
        selected_rows = impact_table.copy()
        
    # Make a copy of selected rows - looks useless but reduces error messages for some reason
    selected_rows = selected_rows.copy()

    # Create empty column for full references
    selected_rows['NAS_Reference'] = None

    # Pull cut-and-paste references from NAS
    for index, row in selected_rows.iterrows():
        url = 'https://nas.er.usgs.gov/queries/references/ReferenceViewer.aspx?refnum=' + str(row['Reference'])
        response = requests.post(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, "html.parser")
            desired_span_tag = soup.find("span", {"id": "ContentPlaceHolder1_CutPasteRef"})
            if desired_span_tag:
                reference = desired_span_tag.get_text(strip = True)
            else:
                reference = 'Error'
            selected_rows.loc[index, 'NAS_Reference'] = reference

    # Formula to format in-text citations
    def format_citation(citation):
        year = re.findall(r'\d{4}', citation)[0]
        match = re.match(r'^(.+?)\.\s\d{4}\.', citation).group(1).strip()
        ref_parts = match.split(', ')
        n_authors = len(ref_parts) - 1
        if n_authors == 1:
            in_text_citation = str(ref_parts[0]) + ' ' + str(year)
        elif n_authors == 2:
            if 'and' in ref_parts[2]:
                second_author = ref_parts[2].split(' ')[-1].strip()
                in_text_citation = str(ref_parts[0]) + ' and ' + str(second_author) + ' ' + str(year)
            else:
                in_text_citation = str(ref_parts[0]) + ' ' + str(year)
        else:
            in_text_citation = str(ref_parts[0]) + ' et al. ' + str(year)
        return(in_text_citation)

    # Create column with in-text citations
    selected_rows['in_text'] = selected_rows['NAS_Reference'].apply(format_citation)

    # Function to combine impact descriptions with in-text citations
    def combine_text(row):
        text = row['Impact Description']
        if row['Impact Description'].endswith('.'):
            row['Impact Description'] = row['Impact Description'][:-1]  
        combined = str(row['Impact Description']) + ' ' + '(' + str(row['in_text']) + ').'
        return combined

    # Apply the function to each row to create a new column
    selected_rows['Impact Description'] = selected_rows.apply(combine_text, axis=1)

    # Reneame columns
    new_column_names = {'Reference': 'RefNum', 'NAS_Reference': 'Reference'}
    selected_rows.rename(columns = new_column_names, inplace = True)
    
    return selected_rows


# Function for reference Word document
def create_references(dataframe):

    # Extract and sort references
    references = sorted(dataframe['Reference'])

    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('References', level = 1)

    # Add references to the document
    for reference in references:
        doc.add_paragraph(reference)
        
    # Convert doc to bytes to make exporting easier
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes


    
# Set layout width to 100%
st.set_page_config(layout="wide")

# Set up main view
st.image('glansisBanner.png')

st.title("GLANSIS Impact Formatting & References")

st.write("""This app allows GLANSIS team members pull and reformat impact descriptions for risk/impact assessments.""")
st.write("""Impact data is scraped from the GLANSIS Great Lakes Impact Info tables.""")
st.write("""'Impact Descriptions' are revised to include an in-text citation. These 'Impact Descriptions' can then be copy and pasted into updated risk/impact assessment drafts. Additional columns with full references and in-text citations are added to data tables. These data tables can be exported as CSV files.""")
st.write("""Full references can downloaded as word document, where they can easily be copy and pasted into risk/impact assessment drafts.""") 

# Set sidebar layout
st.sidebar.title("User Inputs")

st.sidebar.write("""INSTRUCTIONS: To get species specific impacts, users will need to enter the species ID number. If users want only the impacts they have entered for most recently, they will need to enter the impact ID of the first impact they entered. Every subsequent impact ID will be put into a data table. Entering a impact ID is optional.""", wide_mode = True)

species_id = st.sidebar.text_input(f"Species ID Number")
first_impact_id = st.sidebar.text_input(f"Starting Impact ID (optional)")


# Begin data scraping once 'Run' button is clicked
if st.sidebar.button('Run'):
    if species_id:
        with st.spinner('Extracting data...'):
            try:
                df = scrape_impacts(species_id)
                if not df.empty:
                    st.success('Scraping completed!')
                    document = create_references(df)
                else:
                    st.error('Failed to scrape data.')
            except:
                st.error('Failed to scrape data.')
    else:
        st.error('Please enter an Species ID.')
        
        
# Sidebar markdown lines for exporting impact data as CSV
if 'df' in locals():
    st.dataframe(df)
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### Download Data")
    
    csv = df.to_csv(index=False)
    b64_csv = base64.b64encode(csv.encode()).decode()  # Encode CSV data in base64
    href_csv = f'<a href="data:file/csv;base64,{b64_csv}" download="impacts.csv">Download Impacts</a>'
    st.sidebar.markdown(href_csv, unsafe_allow_html =True)

# Sidebar markdown lines for exporting Word Doc with references
if 'document' in locals():
    b64_doc = base64.b64encode(document.read()).decode()  
    href_doc = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_doc}" download="reference.docx">Download References</a>'
    st.sidebar.markdown(href_doc, unsafe_allow_html =True)

    
